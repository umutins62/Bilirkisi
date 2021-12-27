import os
from PyQt5.QtCore import QSettings
from PyQt5.QtWidgets import QWidget, QTableWidget, QApplication, QPushButton, QVBoxLayout, \
    QGridLayout, QGroupBox, \
    QLineEdit, QHeaderView, QTableWidgetItem, QLabel, QAbstractItemView, QMessageBox, QHBoxLayout, QCheckBox, QComboBox
import sys
import sqlite3
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
import settings
import smtplib
from email.message import EmailMessage


class Bilirkisi(QWidget):
    def __init__(self):
        super().__init__()
        self.setUI()
        self.SetUI1()
        self.ayargizle()
        self.getsetting()
        self.temayukle()
        self.kayitgizle()
        self.toplamgizle()
        self.ortalamagizle()
        self.sifregizle()




        # print(self.settings.fileName())
        # textbox
        self.email_value=self.setting_value.value('text box')
        self.sifre_value=self.setting_value1.value('text box1')
        self.gun_value=self.setting_value5.value('text box2')

        self.email.setText(self.email_value)
        self.sifre.setText(self.sifre_value)
        self.gun.setText(self.gun_value)

        # chekbox

        self.sifregoster.stateChanged.connect(self.sifregizle)
        self.sifregoster.stateChanged.connect(self.closeEvent)
        self.otyedekle.stateChanged.connect(self.gizle)
        self.otyedekle.stateChanged.connect(self.closeEvent)
        self.kayıtgoster.stateChanged.connect(self.kayitgizle)
        self.kayıtgoster.stateChanged.connect(self.closeEvent)
        self.ortalamagoster.stateChanged.connect(self.closeEvent)
        self.ortalamagoster.stateChanged.connect(self.ortalamagizle)
        self.toplamgoster.stateChanged.connect(self.toplamgizle)


        self.gizle()

        # combobox
        self.theme_value=self.setting_value5.value('theme')
        self.theme.setCurrentText(self.theme_value)
        # form konumu
        try:
            self.move(self.setting_value9.value('window position'))
        except:
            pass

    def getsetting(self):
        # textbox
        self.setting_value=QSettings('Set_App','vraiables')
        self.setting_value1=QSettings('Set_App','vraiables')
        self.setting_value5=QSettings('Set_App','vraiables')
        # form ayarla
        self.setting_value9=QSettings('Set_App','App1')
        # chekbox
        self.setting_value2=QSettings('A','vraiables')
        self.setting_value4=QSettings('B','vraiables')
        self.setting_value6=QSettings('C','vraiables')
        self.setting_value7=QSettings('D','vraiables')
        self.setting_value10=QSettings('E','vraiables')

        self.sifregoster.setChecked(self.setting_value2.value("A", False, bool))
        self.otyedekle.setChecked(self.setting_value4.value("B", False, bool))
        self.kayıtgoster.setChecked(self.setting_value6.value("C", False, bool))
        self.toplamgoster.setChecked(self.setting_value7.value("D", False, bool))
        self.ortalamagoster.setChecked(self.setting_value10.value("E", False, bool))
        # combobox
        self.setting_value8=QSettings('Set_App','vraiables')

    def closeEvent(self, event):
        # textbox
        self.setting_value.setValue('text box',self.email.text())
        self.setting_value1.setValue('text box1',self.sifre.text())
        self.setting_value5.setValue('text box2',self.gun.text())
        # chekbox
        self.setting_value2.setValue("A", self.sifregoster.isChecked())
        self.setting_value4.setValue("B", self.otyedekle.isChecked())
        self.setting_value6.setValue("C", self.kayıtgoster.isChecked())
        self.setting_value7.setValue("D", self.toplamgoster.isChecked())
        self.setting_value10.setValue("E", self.ortalamagoster.isChecked())
        # combobox
        self.setting_value8.setValue('theme',self.theme.currentText())
        # form ayarla
        self.setting_value9.setValue('window position',self.pos())

    def gizle(self):
        if self.otyedekle.isChecked():
            self.gun.setVisible(True)
        else:
            self.gun.setVisible(False)

    def sifregizle(self):

        if self.sifregoster.isChecked():
            self.sifre.setEchoMode(QLineEdit.Normal)
        else:
            self.sifre.setEchoMode(QLineEdit.Password)

    def setUI(self):
        self.setWindowTitle("Bilirkişi Dosya Takip Uygulaması")

        grb1 = QGroupBox("VERİLER")
        # grb1.setStyleSheet("QGroupBox { border: 1px solid Green;}")
        grb2 = QGroupBox("DOSYA LİSTESİ")
        # grb2.setStyleSheet("QGroupBox { border: 1px solid Green;}")
        grb3 = QGroupBox("")
        grb4 = QGroupBox("")
        self.grb5=QGroupBox("")
        self.grb5.setStyleSheet("QGroupBox { border: 1px solid Green;}")

        grid = QGridLayout()
        grid.setSpacing(10)
        vbox = QVBoxLayout()

        self.kesif_tarih=QLineEdit()
        self.kesif_tarih.setInputMask("99/99/9999")

        self.kesif_konu=QLineEdit()
        self.kesif_konu.setPlaceholderText("Keşif Konusu Giriniz")

        self.kesif_ucreti=QLineEdit()
        self.kesif_ucreti.setPlaceholderText("Keşif Ücretini Giriniz")

        self.kesif_yapan_mahkeme=QLineEdit()
        self.kesif_yapan_mahkeme.setPlaceholderText("Keşif Yapan Mahkemeyi Giriniz")

        self.Teslim_tarih=QLineEdit()
        self.Teslim_tarih.setInputMask("99/99/9999")

        self.kesif_ara = QLineEdit()
        self.kesif_ara.setPlaceholderText("Dosya Ara")
        self.kesif_ara.textChanged.connect(self.Dosya_Ara)

        self.Ekle=QPushButton("Ekle")
        self.Ekle.clicked.connect(self.Dosya_Ekle)
        self.Duzenle=QPushButton("Düzenle")
        self.Duzenle.clicked.connect(self.Dosya_Duzenle)
        self.Sil=QPushButton("Sil")
        self.Sil.clicked.connect(self.Dosya_Sil)
        self.raporla=QPushButton("Rapor")
        self.raporla.clicked.connect(self.Dosya_Raporla)
        self.settins=QCheckBox("Ayarlar")
        self.settins.clicked.connect(self.ayargizle)


        self.L_Tarih=QLabel("Keşif Tarihi")
        self.L_T_Tarih=QLabel("Teslim Tarihi")

        self.L_Hata=QLabel("Uyarılar:")
        self.L_Sayi=QLabel("Toplam Kayıt")
        self.L_Toplam=QLabel("Toplam Alınan Net Tutar")
        self.L_Bos=QLabel("")
        self.L_ortalama=QLabel("Ortalama Rapor Süresi :")

        self.L_Bos1=QLabel("  Bu program umut ÇELİK tarafından tasarlanmıştır.İzinsiz kullanılamaz.")

        self.tableWidget = QTableWidget()
        self.tableWidget.setAlternatingRowColors(True)
        self.tableWidget.setColumnCount(8)

        self.tableWidget.horizontalHeader().setCascadingSectionResizes(False)
        self.tableWidget.horizontalHeader().setSortIndicatorShown(False)
        self.tableWidget.horizontalHeader().setStretchLastSection(True)
        self.tableWidget.verticalHeader().setVisible(False)
        self.tableWidget.verticalHeader().setCascadingSectionResizes(False)
        self.tableWidget.verticalHeader().setStretchLastSection(False)
        self.tableWidget.setHorizontalHeaderLabels(("Sıra No","Tarihi","Konusu", "Mahkeme"
                                                    ,"Ücreti","Net\nKalan", "Teslim\nTarihi","Hazırlanış\nSüresi"))
        self.tableWidget.setEditTriggers(QAbstractItemView.NoEditTriggers)
        style = ":section {""background-color: silver ; }"
        self.tableWidget.horizontalHeader().setStyleSheet(style)
        self.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
        # self.tableWidget.setSelectionMode(QAbstractItemView.MultiSelection)
        self.tableWidget.doubleClicked.connect(self.Listeden_Aktar)

        grid.addWidget(self.L_Tarih, 1, 0)
        grid.addWidget(self.L_T_Tarih, 1, 2)
        grid.addWidget(self.kesif_tarih, 2, 0)
        grid.addWidget(self.kesif_konu, 2, 1)
        grid.addWidget(self.Teslim_tarih, 2, 2)
        grid.addWidget(self.kesif_yapan_mahkeme, 3, 0)
        grid.addWidget(self.kesif_ucreti, 3, 1)

        grid.addWidget(self.Ekle, 2, 4)
        grid.addWidget(self.Duzenle, 3, 4)
        grid.addWidget(self.settins, 1, 5)
        grid.addWidget(self.Sil, 2, 5)
        grid.addWidget(self.raporla, 3, 5)


        grb1.setLayout(grid)
        vbox1 = QVBoxLayout()
        vbox1.addWidget(self.tableWidget)
        grb2.setLayout(vbox1)

        grid1 = QGridLayout()
        grid1.setSpacing(10)
        grb3.setLayout(grid1)

        grid1.addWidget(self.L_Hata, 1, 0)
        grid1.addWidget(self.L_Bos, 1, 1)
        grid1.addWidget(self.L_Bos, 1, 2)
        grid1.addWidget(self.L_ortalama, 1, 3)
        grid1.addWidget(self.L_Sayi, 1, 4)
        grid1.addWidget(self.L_Toplam, 1, 5)

        h2box=QHBoxLayout()
        h2box.addWidget(self.kesif_ara)
        h2box.addStretch()
        grb4.setLayout(h2box)



        vbox.addWidget(grb1)
        vbox.addWidget(grb2)
        vbox.addWidget(grb4)
        vbox.addWidget(grb3)
        vbox.addWidget(self.L_Bos1)
        vbox.addWidget(self.grb5)

        self.setLayout(vbox)
        self.create_connection()
        self.loaddata()
        self.show()

    def SetUI1(self):
        self.setWindowTitle("Ayarlar")
        hbox=QHBoxLayout()
        grub1=QGroupBox("")
        grub2=QGroupBox("")
        self.dataadet=QLabel("Kayıt Sayısını Göster")
        self.datatoplam=QLabel("Toplam Net Tutarı  Göster")
        self.ortgoster=QLabel("Ortalama Cevap Süresi Göster")

        self.theme=QComboBox()
        self.theme.addItem("Light")
        self.theme.addItem("Consol Style")
        self.theme.addItem("ElegantDark")
        self.theme.addItem("Ubuntu")
        self.theme.addItem("MacOS")
        self.theme.addItem("Material Dark")
        self.theme.addItem("Roblabla")

        self.theme.currentTextChanged.connect(self.temayukle)
        self.tema=QLabel("Tema")


        #veritabanı ayarları
        self.email=QLineEdit()
        self.email.setPlaceholderText("Email Adresini Gir")
        self.sifre=QLineEdit()
        self.sifre.setPlaceholderText("Şifreyi Gir")
        self.sil=QPushButton("Veritabanını Sil")
        self.sil.clicked.connect(self.veritabanisil)
        self.yedekle=QPushButton("Veritabanını Yedekle")
        self.yedekle.clicked.connect(self.veritabaniyedekle)

        self.sifregoster=QCheckBox("Şifreyi Göster")
        self.kayıtgoster=QCheckBox("")
        self.toplamgoster=QCheckBox("")
        self.ortalamagoster=QCheckBox("")


        self.otyedekle=QCheckBox("Otomatik Yedekle(Gün)")
        self.gun=QLineEdit()
        self.gun.setPlaceholderText("Gün Gir")


        grid1 = QGridLayout()
        grid2 = QGridLayout()
        grub1.setLayout(grid1)
        grub2.setLayout(grid2)

        grid1.addWidget(self.email,1,0)
        grid1.addWidget(self.sifre,2,0)
        grid1.addWidget(self.sifregoster,2,1)
        grid1.addWidget(self.otyedekle,3,0)
        grid1.addWidget(self.gun,3,1)
        grid1.addWidget(self.sil,4,0)
        grid1.addWidget(self.yedekle,4,1)

        grid2.addWidget(self.dataadet,1,0)
        grid2.addWidget(self.kayıtgoster,1,1)
        grid2.addWidget(self.datatoplam,2,0)
        grid2.addWidget(self.toplamgoster,2,1)
        grid2.addWidget(self.ortgoster,3,0)
        grid2.addWidget(self.ortalamagoster,3,1)
        grid2.addWidget(self.tema,4,0)
        grid2.addWidget(self.theme,4,1)


        hbox.addWidget(grub1)
        hbox.addWidget(grub2)
        self.grb5.setLayout(hbox)

    def temayukle(self):
        if self.theme.currentText()=="Consol Style":
            self.setStyleSheet(open('consolstyle.qss','r').read())
        elif self.theme.currentText()=="ElegantDark":
            self.setStyleSheet(open('elegantdark.qss','r').read())
        elif self.theme.currentText()=="Ubuntu":
            self.setStyleSheet(open('ubuntu.qss','r').read())
        elif self.theme.currentText()=="MacOS":
            self.setStyleSheet(open('macos.qss','r').read())
        elif self.theme.currentText()=="Material Dark":
            self.setStyleSheet(open('materialdark.qss','r').read())
        elif self.theme.currentText()=="Roblabla":
            self.setStyleSheet(open('roblabla.qss','r').read())
        elif self.theme.currentText()=="Light":
            self.setStyleSheet(open('Light.qss','r').read())
        else:
            pass

    def ac_settings(self):
        self.sett=settings.ayarlar()

    def create_connection(self):
        if os.path.exists("bilirkisidb.db"):
            print("Veritabanı Var")
        else:
            # Veritabanı Oluştur
            connection_obj = sqlite3.connect('bilirkisidb.db')
            # cursor
            cursor_obj = connection_obj.cursor()
            cursor_obj.execute("DROP TABLE IF EXISTS BILIRKISI")
            # Tabloyu oluştur
            table = """ CREATE TABLE BILIRKISI (
                        K_İd INTEGER  PRIMARY KEY,
                        K_Tarih VARCHAR(255) NOT NULL,
                        K_Konu VARCHAR(255) NOT NULL,
                        K_Mahkeme VARCHAR(255) NOT NULL,
                        K_Ucret REAL NOT NULL,
                        K_Net REAL,
                        K_Teslim VARCHAR(255),  
                        K_Sure INTEGER   
                    ); """
            cursor_obj.execute(table)
            self.L_Hata.setText("Uyarılar : Veritabanı Oluşturuldu")
            self.L_Hata.setStyleSheet("color:blue;font-weight:bold")

    def loaddata(self):
        self.connection = sqlite3.connect("bilirkisidb.db")
        query = "SELECT * FROM BILIRKISI"
        result = self.connection.execute(query)
        self.tableWidget.setRowCount(0)
        for row_number, row_data in enumerate(result):
            self.tableWidget.insertRow(row_number)
            for column_number, data in enumerate(row_data):
                self.tableWidget.setItem(row_number, column_number,QTableWidgetItem(str(data)))

        self.c = self.connection.cursor()

        TOPP = self.c.execute("SELECT COUNT(*) FROM BILIRKISI").fetchall()
        kayit=int(str(TOPP).strip("[](),"))

        if kayit>0:
            reques_ = self.c.execute("SELECT SUM(K_Net) FROM BILIRKISI").fetchall()
            reques_1 = self.c.execute("SELECT SUM(K_Sure) FROM BILIRKISI").fetchall()
            toplam=float(str(reques_).strip("[](),"))

            self.L_Toplam.setText("Toplam Alınan Net Tutar : " + str(round(toplam,2)) + " TL      ")
            self.L_Sayi.setText("Toplam Kayıt : " + str(TOPP).strip("[](),") + " ad       ")
            ortalama=float(str(reques_1).strip("[](),"))/float(str(TOPP).strip("[](),"))
            self.L_ortalama.setText("Ort. Rapor Cevap Süresi : " + str(round(ortalama,2)) + " gün       ")
        else:
            self.L_Toplam.setText("Toplam Alınan Net Tutar : 0.00  TL      ")
            self.L_Sayi.setText("Toplam Kayıt : 0 ad       ")
            self.L_ortalama.setText("Ort. Rapor Cevap Süresi : 0 gün       ")

        self.connection.close()

    def Dosya_Ekle(self):

        K_Tarih=""
        K_Konu=""
        K_Mahkeme=""
        K_Ucret=0
        K_Net=0
        K_Teslim=""
        K_Sure=0
        Damga_Vergisi=0
        Gelir_Vergisi=0


        K_Tarih=self.kesif_tarih.text()
        K_Konu=self.kesif_konu.text()
        K_Mahkeme=self.kesif_yapan_mahkeme.text()
        K_Ucret=float(self.kesif_ucreti.text().replace(",","."))
        Damga_Vergisi=0.15*K_Ucret
        Gelir_Vergisi=7.59/1000*K_Ucret
        K_Net=K_Ucret-(Damga_Vergisi+Gelir_Vergisi)
        K_Teslim=self.Teslim_tarih.text()



        month = K_Teslim
        date = datetime.date(int(month.split('/')[2]),int(month.split('/')[1]),int( month.split('/')[0]))
        print(date)

        month1 = K_Tarih
        date1 = datetime.date(int(month1.split('/')[2]),int(month1.split('/')[1]),int( month1.split('/')[0]))
        print(date1)

        fark=date-date1
        K_Sure=int(fark.days)

        print(Damga_Vergisi)
        print(Gelir_Vergisi)
        print(str(K_Net))
        print(str(K_Sure))


        try:

            self.conn = sqlite3.connect("bilirkisidb.db")
            self.c = self.conn.cursor()
            self.c.execute("INSERT INTO BILIRKISI (K_Tarih,K_Konu,K_Mahkeme,K_Ucret,K_Net,K_Teslim,K_Sure) VALUES (?,?,?,?,?,?,?)",(K_Tarih,K_Konu,K_Mahkeme,K_Ucret,K_Net,K_Teslim,K_Sure))

            self.conn.commit()
            self.c.close()
            self.conn.close()
            self.L_Hata.setText("Uyarılar : Kayıt başarılı")
            self.L_Hata.setStyleSheet("color:blue;font-weight:bold")
            self.loaddata()




        except:
            self.L_Hata.setText("Uyarılar : Kayıt yapılırken hata oluştu!")
            self.L_Hata.setStyleSheet("color:blue;font-weight:bold")

    def Listeden_Aktar(self,item):
        sf = "You clicked on {0}x{1}".format(item.column(), item.row())
        # print(sf)
        global book_id
        book_list=[]
        for i in range(0,7):
            book_list.append(self.tableWidget.item(self.tableWidget.currentRow(),i).text())

        book_id=book_list[0]
        self.kesif_tarih.setText(book_list[1])
        self.kesif_konu.setText(book_list[2])
        self.kesif_yapan_mahkeme.setText(book_list[3])
        self.kesif_ucreti.setText(book_list[4])
        self.Teslim_tarih.setText(book_list[6])
        print(book_id)

    def Dosya_Duzenle(self):
        K_Tarih=self.kesif_tarih.text()
        K_Konu=self.kesif_konu.text()
        K_Mahkeme=self.kesif_yapan_mahkeme.text()
        K_Ucret=float(self.kesif_ucreti.text().replace(",","."))
        Damga_Vergisi=0.15*K_Ucret
        Gelir_Vergisi=7.59/1000*K_Ucret
        K_Net=K_Ucret-(Damga_Vergisi+Gelir_Vergisi)
        K_Teslim=self.Teslim_tarih.text()



        month = K_Teslim
        date = datetime.date(int(month.split('/')[2]),int(month.split('/')[1]),int( month.split('/')[0]))
        print(date)

        month1 = K_Tarih
        date1 = datetime.date(int(month1.split('/')[2]),int(month1.split('/')[1]),int( month1.split('/')[0]))
        print(date1)

        fark=date-date1
        K_Sure=int(fark.days)

        try:
            self.conn = sqlite3.connect("bilirkisidb.db")
            self.c = self.conn.cursor()
            self.c.execute("update BILIRKISI set K_Tarih=?,K_Konu=?,K_Mahkeme=?,K_Ucret=?,K_Net=?,K_Teslim=?,K_Sure=? where K_İd=?",(K_Tarih,K_Konu,K_Mahkeme,K_Ucret,K_Net,K_Teslim,K_Sure,book_id))
            self.conn.commit()
            self.c.close()
            self.conn.close()
            self.L_Hata.setText("Uyarılar : Güncelleme başarılı")
            self.L_Hata.setStyleSheet("color:blue;font-weight:bold")
            self.loaddata()




        except:
            self.L_Hata.setText("Uyarılar : Güncelleme yapılırken hata oluştu!")
            self.L_Hata.setStyleSheet("color:blue;font-weight:bold")

    def Dosya_Sil(self):
        onay=QMessageBox.question(self,"Uyarı","Silmek istediğinizden eminmisiniz?",QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

        if (onay==QMessageBox.Yes):
            try:
                self.conn = sqlite3.connect("bilirkisidb.db")
                self.c = self.conn.cursor()
                self.c.execute("delete from BILIRKISI where K_İd=?",(book_id,))
                self.conn.commit()
                self.c.close()
                self.conn.close()
                self.loaddata()

                self.L_Hata.setText("Uyarılar : Silme işlemi başarılı")
                self.L_Hata.setStyleSheet("color:blue;font-weight:bold")


            except:
                self.L_Hata.setText("Uyarılar : Silme işlemi yapılırken hata oluştu!")
                self.L_Hata.setStyleSheet("color:blue;font-weight:bold")

    def Dosya_Ara(self):
        try:
            value = self.kesif_ara.text()
            if value == "":
                self.loaddata()
                pass
            else:
                self.conn = sqlite3.connect("bilirkisidb.db")
                self.c = self.conn.cursor()
                query = self.c.execute(
                    "SELECT K_İd,K_Tarih,K_Konu,K_Mahkeme,K_Ucret,K_Net,K_Teslim,K_Sure FROM BILIRKISI WHERE K_Konu LIKE ?",
                    ('%' + value + '%',)).fetchall()
                # print(query)
                if query == []:
                    QMessageBox.information(self, "Dikkat", "Aradığınız rapor bulunamadı!")
                else:
                    for i in reversed(range(self.tableWidget.rowCount())):
                        self.tableWidget.removeRow(i)
                    for row_data in query:
                        row_number = self.tableWidget.rowCount()
                        self.tableWidget.insertRow(row_number)

                        for colum_header, data in enumerate(row_data):
                            self.tableWidget.setItem(row_number, colum_header, QTableWidgetItem(str(data)))
        except:
            pass

    # def Veritabani_Bilgileri(self):
    #
    #     self.conn = sqlite3.connect("bilirkisidb.db")
    #     self.c = self.conn.cursor()
    #     reques_ = self.c.execute("SELECT SUM(K_Net) FROM BILIRKISI").fetchall()
    #     reques_1 = self.c.execute("SELECT SUM(K_Sure) FROM BILIRKISI").fetchall()
    #     TOPP = self.c.execute("SELECT COUNT(*) FROM BILIRKISI").fetchall()
    #     self.L_Toplam.setText("Toplam Alınan Net Tutar : " + str(reques_).strip("[](),") + " TL      ")
    #     self.L_Sayi.setText("Toplam Kayıt : " + str(TOPP).strip("[](),") + " ad       ")
    #     ortalama=float(str(reques_1).strip("[](),"))/float(str(TOPP).strip("[](),"))
    #     self.L_ortalama.setText("Ort. Rapor Cevap Süresi : " + str(ortalama) + " gün       ")
    #
    #     self.conn.commit()
    #     self.c.close()
    #     self.conn.close()
    #     self.loaddata()

    def Dosya_Raporla(self):

        self.conn = sqlite3.connect("bilirkisidb.db")
        self.c = self.conn.cursor()
        reques_ = self.c.execute("SELECT SUM(K_Net) FROM BILIRKISI").fetchall()
        TOPP = self.c.execute("SELECT COUNT(*) FROM BILIRKISI").fetchall()
        self.L_Toplam.setText("Toplam Alınan Net Tutar : " + str(reques_).strip("[](),") + " TL      ")
        self.L_Sayi.setText("Toplam Kayıt : " + str(TOPP).strip("[](),") + " ad       ")
        # salary_list=[]
        # for i in reques_:
        #     salary_list.append(i)
        # print(salary_list)
        self.conn.commit()
        self.c.close()
        self.conn.close()





        document = Document()
        logo_path = 'dsi.png'  # Path of the image file
        section = document.sections[-1]  # Create a section
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        sec_header = section.header  # Create header
        header_tp = sec_header.add_paragraph()  # Add a paragraph in the header, you can add any anything in the paragraph
        header_run = header_tp.add_run()  # Add a run in the paragraph. In the run you can set the values
        header_run.add_picture(logo_path, width=Inches(0.6))  # Add a picture and set width.
        rml_header = "\t\t DSİ 93. ŞUBE MÜDÜRLÜĞÜ  \t"
        header_run.add_text(rml_header)
        header_tp.alignment = 0

        document.add_heading('BİLİRKİŞİLİK DOSYA LİSTESİ', level=2)
        paragraph = document.add_paragraph()
        paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

        table1 = document.add_table(rows=2, cols=2)


        heading_cells = table1.rows[0].cells
        heading_cells[0].text = 'Alınan Toplam Net Tutar : '+str(reques_).strip("[](),")+" TL"

        heading_cells = table1.rows[1].cells
        heading_cells[0].text = 'Toplam Kayıt Sayısı : '+str(TOPP).strip("[](),")+" ad"






        self.conn = sqlite3.connect("bilirkisidb.db")
        self.c = self.conn.cursor()
        records1 = self.c.execute(
            "SELECT K_Tarih, K_Konu, K_Mahkeme, K_Ucret, K_Net, K_Teslim, K_Sure FROM BILIRKISI").fetchall()
        sayi = self.c.execute('SELECT COUNT(*) from BILIRKISI')
        row = self.c.fetchone()


        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False
        for cell in table.rows[0].cells:
            cell.width = Inches(0.7)

        hdr_cells = table.rows[0].cells
        # K_Tarih, K_Konu, K_Mahkeme, K_Ucret, K_Net, K_Teslim, K_Sure

        hdr_cells[0].text = 'Keşif Tarihi'
        hdr_cells[1].text = 'Keşif Konusu'
        hdr_cells[2].text = 'Keşif Mahkeme'
        hdr_cells[3].text = 'Keşif Ücreti'
        hdr_cells[4].text = 'Keşif Net Kalan Tutar'
        hdr_cells[5].text = 'Keşif Teslim Tarihi'
        hdr_cells[6].text = 'Rapor Teslim Süresi'




        for K_Tarih, K_Konu, K_Mahkeme, K_Ucret, K_Net, K_Teslim, K_Sure in records1:
            row_cells = table.add_row().cells

            row_cells[0].text = K_Tarih
            row_cells[1].text = K_Konu
            row_cells[2].text = K_Mahkeme
            row_cells[3].text = str(K_Ucret)
            row_cells[4].text = str(K_Net)
            row_cells[5].text = K_Teslim
            row_cells[6].text = str(K_Sure)



        # document.add_page_break()

        document.save('Bilirkisilik_Dosyalari.docx')

        os.system("start Bilirkisilik_Dosyalari.docx")

    def ayargizle(self):
        if self.settins.isChecked():
            self.grb5.setVisible(True)
        else:
            self.grb5.setVisible(False)

    def kayitgizle(self):

        if self.kayıtgoster.isChecked():
            self.L_Sayi.setVisible(True)
        else:
            self.L_Sayi.setVisible(False)

    def toplamgizle(self):


        if self.toplamgoster.isChecked():
            self.L_Toplam.setVisible(True)
        else:
            self.L_Toplam.setVisible(False)

    def ortalamagizle(self):


        if self.ortalamagoster.isChecked():
            self.L_ortalama.setVisible(True)
        else:
            self.L_ortalama.setVisible(False)

    def veritabanisil(self):
        onay=QMessageBox.question(self,"Uyarı","Veritabanını silmek istediğinizden eminmisiniz?",QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

        if (onay==QMessageBox.Yes):
            try:

                os.remove("bilirkisidb.db")
                self.L_Hata.setText("Uyarılar : Veritabanı silme işlemi başarılı")
                self.L_Hata.setStyleSheet("color:blue;font-weight:bold")

                onay1 = QMessageBox.question(self, "Uyarı", "Veritabanını oluşturulsunmu?",
                                            QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

                if (onay1 == QMessageBox.Yes):
                    self.create_connection()
                    self.loaddata()
                else:
                    pass

            except:
                self.L_Hata.setText("Uyarılar : Veritabanı silme işlemi yapılırken hata oluştu!")
                self.L_Hata.setStyleSheet("color:blue;font-weight:bold")

    def veritabaniyedekle(self):
        onay=QMessageBox.question(self,"Uyarı","Veritabanını yedeklemek istediğinizden eminmisiniz?",QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

        if (onay==QMessageBox.Yes):
            try:


                EMAIL_ADDRESS = self.email.text()
                EMAIL_PASSWORD = self.sifre.text()
                msg = EmailMessage()
                msg['Subject'] = 'Veritaabnı yedekleme'
                msg['From'] = EMAIL_ADDRESS
                msg['To'] = EMAIL_ADDRESS
                msg.set_content('Veritabanı yedeklendi...')

                with open('bilirkisidb.db', 'rb') as f:
                    file_data = f.read()
                    file_name = f.name

                msg.add_attachment(file_data, maintype='_application_', subtype='octet-stream', filename=file_name)
                with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
                    smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
                    smtp.send_message(msg)

                self.L_Hata.setText("Uyarılar : Yedekleme işlemi başarılı")
                self.L_Hata.setStyleSheet("color:blue;font-weight:bold")



            except:
                self.L_Hata.setText("Uyarılar : .Yedekleme işlemi yapılırken hata oluştu!")
                self.L_Hata.setStyleSheet("color:blue;font-weight:bold")




if __name__ == "__main__":
    app = QApplication(sys.argv)
    pencere = Bilirkisi()
    sys.exit(app.exec())